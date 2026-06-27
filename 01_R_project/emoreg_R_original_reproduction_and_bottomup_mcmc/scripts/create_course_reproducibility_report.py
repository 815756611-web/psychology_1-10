"""Build the standalone course reproducibility report from the appendix template."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT.parents[1] / "03_paper_and_course_materials" / "心理学院R编程语言课程可重复检验指南(2025版).docx"
OUT = ROOT / "docs" / "Bo_2024_情绪调节计算可复现性检验报告.docx"


def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph, text, size=11, bold=False, align=None, line_spacing=1.5):
    clear_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = line_spacing
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def checkbox_block(options, selected):  # options: 备选项列表；selected: 选中的精确文本；作用：把模板里的“是/否”类单选项改成勾选式显示。
    lines = []
    for option in options:
        marker = "☑" if option == selected else "☐"
        lines.append(f"{marker} {option}")
    return "\n".join(lines)


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_paragraph_index(doc, text):
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return index
    raise ValueError(f"Paragraph index not found: {text}")


def prune_before_title(doc):
    title = find_paragraph(doc, "对作者(年份)研究结果的计算可复现性检验")
    body = doc._element.body
    while body[0] != title._element:
        body.remove(body[0])


def fill_group_table(table):
    set_cell_text(table.cell(0, 1), "本次提交者（个人完成）", size=11)
    set_cell_text(table.cell(1, 1), "本次提交者（个人完成）", size=11)
    set_cell_text(table.cell(3, 1), "本次提交者：数据整理、R 脚本运行、结果核对", size=10.5)
    set_cell_text(table.cell(3, 3), "本次提交者：PPT 图表整理", size=10.5)
    set_cell_text(table.cell(4, 1), "本次提交者：报告撰写", size=10.5)
    set_cell_text(table.cell(4, 3), "本次提交者：PPT 展示", size=10.5)


def fill_literature_table(table):
    set_cell_text(
        table.cell(1, 1),
        "Bo, K., Kraynak, T. E., Kwon, M., Sun, M., Gianaros, P. J., & Wager, T. D. "
        "(2024). A systems identification approach using Bayes factors to deconstruct "
        "the brain bases of emotion regulation. Nature Neuroscience. "
        "https://doi.org/10.1038/s41593-024-01605-7",
        size=10,
    )
    set_cell_text(
        table.cell(2, 1),
        "NeuroVault collection 16266 的一阶 beta 图与行为数据；原作者 MATLAB/CANlab "
        "工作流备份；副本 R 项目 emoreg_R_original_reproduction_and_bottomup_mcmc。",
        size=10,
    )
    set_cell_text(
        table.cell(4, 1),
        checkbox_block(
            ["是，且包含元分析研究", "是，但不包含元分析研究", "否"],
            "是，但不包含元分析研究"
        ),
        size=10,
    )
    set_cell_text(
        table.cell(4, 3),
        checkbox_block(
            ["是(附上原文链接)", "否"],
            "否"
        ),
        size=10,
    )
    set_cell_text(
        table.cell(6, 1),
        "Route A：重复原文两类关键行为效应，并按 BF/t 逻辑识别四类脑系统图，再比较 AHAB/PIP 共识图与原文图。"
        "Route B：把原文四系统图拆成 61 个 ROI 连通簇，在 AHAB 与 PIP 中分别进行 ROI 层级贝叶斯 MCMC 聚类，并比较 K=2~6 的后验结构。",
        size=10,
    )
    set_cell_text(
        table.cell(7, 1),
        checkbox_block(
            ["是(附上原文链接)", "否"],
            "否"
        ),
        size=10,
    )
    set_cell_text(
        table.cell(7, 3),
        "原文主体分析依托 2 个独立样本（AHAB、PIP）。Route A 对两者分别复现并做跨样本共识图；Route B 也在两样本中分别建模后，再比较 cluster 的跨样本稳定性。",
        size=10.5,
    )
    set_cell_text(
        table.cell(8, 1),
        "因为 Route A 可以直接检验原文主结论能否在公开材料上被完整重算，Route B 则进一步检验这些理论系统在数据驱动后验结构中是否稳定出现；"
        "两条路线回答的问题不同，但互补。",
        size=10,
    )
    set_cell_text(
        table.cell(10, 1),
        checkbox_block(
            ["是", "否"],
            "是"
        ) + "\n采用作者公开的一阶 beta 图和行为数据。",
        size=10,
    )
    set_cell_text(
        table.cell(10, 3),
        checkbox_block(
            ["是(说明原因)", "否"],
            "否"
        ),
        size=10,
    )
    set_cell_text(table.cell(11, 1), "原文 AHAB n=182、PIP n=176；本次使用相同样本量。", size=10.5)
    set_cell_text(table.cell(11, 3), "不适用，本次未修改样本量。", size=10.5)


def fill_descriptive_table(table):
    set_cell_text(table.cell(0, 1), "Route A：情绪反应性\nLook negative - Look neutral", size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 4), "Route A：重评效应\nRegulate negative - Look negative", size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)

    values = {
        (2, 1): "358",
        (2, 2): "NR",
        (2, 3): "NR",
        (2, 4): "358",
        (2, 5): "NR",
        (2, 6): "NR",
        (3, 1): "358",
        (3, 2): "2.360",
        (3, 3): "0.700",
        (3, 4): "358",
        (3, 5): "-0.608",
        (3, 6): "0.659",
        (4, 1): "0%",
        (4, 2): "NR",
        (4, 3): "NR",
        (4, 4): "0%",
        (4, 5): "NR",
        (4, 6): "NR",
        (5, 1): "完全一致",
        (5, 2): "原文未报告",
        (5, 3): "原文未报告",
        (5, 4): "完全一致",
        (5, 5): "原文未报告",
        (5, 6): "原文未报告",
    }
    for (row, col), value in values.items():
        set_cell_text(table.cell(row, col), value, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def fill_original_inference_table(table):
    set_cell_text(table.cell(0, 1), "Route A：关键行为效应", size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 6), "Route A：脑系统空间复现", size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 6), "共识图指标", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    values = {
        (2, 1): "358\n358",
        (2, 2): "63.8\n-17.5",
        (2, 3): "3.37\n0.92",
        (2, 4): "< .001\n< .001",
        (2, 5): "情绪生成为正\n重评效应为负",
        (2, 6): "原文为 4 类二值系统图",
        (3, 1): "358\n358",
        (3, 2): "63.736\n-17.453",
        (3, 3): "3.369\n0.922",
        (3, 4): "< .001\n< .001",
        (3, 5): "与原文方向一致",
        (3, 6): "覆盖率 0.977-0.999\nDice 0.693-0.827",
        (4, 1): "0%\n0%",
        (4, 2): "0.10%\n0.27%",
        (4, 3): "0.04%\n0.26%",
        (4, 4): "不适用\n不适用",
        (4, 5): "均为舍入级差异",
        (4, 6): "不使用 PE，改用空间重叠",
        (5, 1): "完全一致",
        (5, 2): "因舍入导致的偏差",
        (5, 3): "因舍入导致的偏差",
        (5, 4): "一致",
        (5, 5): "推论一致",
        (5, 6): "高覆盖，但边界更宽",
    }
    for (row, col), value in values.items():
        set_cell_text(table.cell(row, col), value, size=9.8, align=WD_ALIGN_PARAGRAPH.CENTER)


def fill_innovation_table(table):
    set_cell_text(table.cell(0, 0), "Route B 项目", size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 0), "AHAB：后验 K 比较", size=10.2)
    set_cell_text(table.cell(2, 0), "PIP：后验 K 比较", size=10.2)
    set_cell_text(table.cell(3, 0), "跨样本：cluster 稳定性", size=10.2)

    rows = {
        1: ["61 ROI（182 subjects）", "overall_recommended_K=3；overall_score=0.8912", "occupied_K=2；effective_K=2；mean max posterior=1.000", "不适用", "主类为 common-appraisal-like；次类更接近 modifiable-emotion-like"],
        2: ["61 ROI（176 subjects）", "overall_recommended_K=2；overall_score=0.9779", "occupied_K=2；effective_K=2；mean max posterior=1.000", "不适用", "主类为 common-appraisal-like；次类更接近 emotion-related mixed 类"],
        3: ["61 shared ROI", "ARI=0.6253；NMI=0.4180", "matched cluster Dice=0.9254 / 0.6860", "不适用", "large common-appraisal-like 类别跨样本稳定；小类一致性较弱"],
    }
    for row, values in rows.items():
        for offset, value in enumerate(values, start=1):
            align = WD_ALIGN_PARAGRAPH.CENTER if offset < 5 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(row, offset), value, size=9.8, align=align)


def fill_reproducibility_tables(table5, table6, table7):
    counts = [
        ("完全一致(δ = 0%)", "2", "33.3%"),
        ("偏差较小(0% < δ < 10%)", "0", "0%"),
        ("偏差较大(δ ＞ 10%)", "0", "0%"),
        ("因舍入导致的偏差", "4", "66.7%"),
    ]
    for row, (label, n_value, pct_value) in enumerate(counts, start=2):
        set_cell_text(table5.cell(row, 0), label, size=10)
        set_cell_text(table5.cell(row, 1), n_value, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table5.cell(row, 2), pct_value, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(table6.cell(2, 1), "6", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table6.cell(2, 2), "100%", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table6.cell(3, 1), "0", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table6.cell(3, 2), "0%", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(table7.cell(2, 1), "1", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table7.cell(2, 2), "33.3%", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table7.cell(3, 1), "2", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table7.cell(3, 2), "66.7%", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def fill_reason_table(table):
    set_cell_text(table.cell(0, 3), "本研究（Route A/B 综合）", size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 4), "", size=10)
    set_cell_text(table.cell(0, 5), "", size=10)

    judgments = {
        1: "否",
        2: "否",
        3: "部分",
        4: "否",
        5: "部分",
        6: "否",
        7: "否",
        8: "是",
        9: "否",
        10: "部分",
        11: "否",
        12: "是",
        13: "是",
        14: "否",
        15: "部分",
        16: "不作为主要原因",
        17: "否",
        18: "否",
        19: "不作为主要原因",
        20: "未在本次材料中确认",
        21: "未在本次材料中确认",
    }
    for row, value in judgments.items():
        set_cell_text(table.cell(row, 3), value, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(row, 4), "", size=9.5)
        set_cell_text(table.cell(row, 5), "", size=9.5)


def fill_paragraphs(doc):
    set_paragraph_text(
        find_paragraph(doc, "对作者(年份)研究结果的计算可复现性检验"),
        "对 Bo 等（2024）研究结果的计算可复现性检验",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_paragraph_text(
        find_paragraph(doc, "摘要：200 ～ 400字，包括背景、方法、结果与讨论/结论几个关键部分"),
        "摘要：本研究将对 Bo 等（2024）的计算可复现性检验明确分成两条路线。Route A 严格采用原文思路，在 AHAB（n=182）和 PIP（n=176）两个独立样本中重建 Look negative - Look neutral 与 Regulate negative - Look negative 两个核心对比，使用 JZS Bayes factor 与原文同阈值逻辑识别四类情绪调节系统，并比较 AHAB/PIP 共识图与原文图的空间重叠。结果表明，两个关键行为效应的 t 值与效应量均复现到原文舍入精度；四类系统对原文图的覆盖率几乎为 1，Dice 为 0.693 至 0.827，主要差异表现为边界外扩。Route B 则把原文四系统图拆成 61 个 ROI 连通簇，在 K=2 至 K=6 之间比较贝叶斯层级混合聚类模型，并用后验模型表现、归属稳定性、共聚类结构和可解释性综合推荐 cluster 数。Route B 支持一个跨样本稳定的大型 common-appraisal-like 类别和一个较小的情绪相关类别，但未恢复四个稳定且非空的终端 cluster。总体上，Route A 支持原文核心结论，Route B 则提示原文四系统更接近理论规则划分，而不是后验自动聚类会自然恢复的四类结构。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "关键词：三到五个关键词，需要包括“计算可复现性”"),
        "关键词：计算可复现性；认知重评；情绪调节；Bayes factor；fMRI",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "简要介绍所选文献的研究背景、主要研究问题及假设、研究结果和结论"),
        "Bo 等（2024）关注一个经典但长期混杂的问题：当个体对负性刺激进行认知重评时，脑内哪些区域属于重评本身，哪些区域反映负性评价的共同过程，哪些区域虽参与情绪生成却不被重评改变。为避免仅凭显著性差异做事后解释，原文以 Bayes factor 同时评估“存在效应”和“支持零效应”的证据，在两个独立样本中识别出四类系统。作者报告负性图片显著提高负性情绪、重评显著降低负性情绪，并认为四类系统在两个样本中均可复制。本次报告围绕这些主结论进行计算可复现性检验。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "对原文献的数据集进行说明。"),
        "本次复现沿用作者公开材料中可直接用于主分析的全部样本：AHAB 182 人，PIP 176 人，合计 358 人。每名被试均具有 Look neutral、Look negative 和 Regulate negative 条件的一阶 beta 图，以及相应的情绪评分数据。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "如使用的新样本与原文样本不一致，还需补充理由，并报告具体的改动标准、步骤，新样本信息和改动后的效应量。"),
        "本次未删改样本量，也未从原始 BOLD 序列重新预处理，因此检验目标是“公开分析输入是否可被重算”，不是“从原始扫描到最终图像的全流程再现”。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "介绍研究设计、数据分析的方法，使用的软件及软件包。"),
        "本次方法部分按路线拆开报告。Route A 沿用原文主分析：对每个体素重算 Look negative - Look neutral（情绪生成）与 Regulate negative - Look negative（重评效应）两个单样本 t 检验，并用与原文一致的 JZS Bayes factor 逻辑判断证据方向：BF > 10 支持有效应，BF < 0.1 支持零效应；在本项目保存的 2log(BF) 标度上对应 +4.6 与 -4.6。随后结合效应方向与基线约束，识别 reappraisal only、common appraisal、non-modifiable emotion generation 和 modifiable emotion generation 四类系统，并按原始工作流保留不少于 15 个体素的簇，再用平滑乘积法生成 AHAB/PIP 共识图。Route B 不再逐体素比较，而是把原文四系统图拆成 61 个 ROI 连通簇，并为每名被试、每个 ROI 提取 emotion_generation、reappraisal_effect、look_neg_base 和 reg_neg_base 四个特征，再在 K=2~6 间比较 ROI 层级贝叶斯 MCMC 混合模型。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "对选用的文章进行复现的思路，包括选取了哪些部分重复及其原因，在使用原文献方法进行重复时是否有需要改动的情况、具体如何改动的。"),
        "Route A 严格围绕原文主结果展开：重算两类关键行为效应、四类体素系统图，以及跨样本共识图与原文图的空间重叠。由于 Windows 下原 BayesFactor 依赖会出现本地退出问题，本次在 R 中实现了与原先验相同的 JZS 数值积分函数，但模型含义与阈值标准不变。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "此外，如果使用了创新方法，需要具体介绍，并结合相应理论对该方法的可行性、及其相比原方法的优势进行说明。"),
        "Route B 是探索性扩展，不替代 Route A。它把原文四系统图中的 61 个 ROI 连通簇视为终端分析单位，保留被试层观测，在 K=2 至 K=6 之间比较贝叶斯层级混合聚类模型。模型通过 MCMC 抽样获得 cluster 归属的后验分布，因此不仅能比较模型表现，还能检查后验归属稳定性、共聚类结构和跨样本一致性。这一路线的优势是能直接评估“数据驱动类别数是否支持原文四系统框架”，代价是它回答的问题与 Route A 并不完全相同。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "对原文献描述性统计进行重复的结果，并汇总表格："),
        "表 2 只对应 Route A，因为 Route A 直接重算原文行为对比；Route B 的输入是被试×ROI 特征和后验模型比较，不使用这一类样本均值/标准差表。原文正文没有给出两个关键行为对比可直接逐项比对的 Mean 和 SD，因此表 2 对原文未报告部分记为 NR，同时报告公开数据复算得到的描述性统计。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "报告对原文献推断性统计进行重复的结果，并汇总表格："),
        "表 3 单独报告 Route A。两个关键行为效应均在公开数据上被精确复算到原文的舍入精度：情绪生成效应 t(357)=63.736，d=3.369；重评效应 t(357)=-17.453，|d|=0.922。与此同时，四类系统的空间复现不再用单一 PE 表示，而改用覆盖率、Dice 和 Jaccard 描述。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "* 此前描述性统计已报告，因此此处仅需进行报告，不需作为一个额外的结果"),
        "Route A 的脑图结果不适合按单一百分误差处理，因此本报告将四类系统的复现评价转为共识图覆盖率、Dice 与 Jaccard。四类系统对原文图的覆盖率均接近 1，说明主要差异不是漏检原文区域，而是共识图边缘更宽；其中 non-modifiable emotion generation 的外扩最明显。",
        size=10.5,
    )
    set_paragraph_text(
        find_paragraph(doc, "对采用新方法进行数据分析的描述性或推断性结果进行描述，并记录在表格中。"),
        "表 4 单独报告 Route B。MCMC 结果显示：AHAB 的 overall_recommended_K 为 3，但实际只形成 2 个非空 cluster；PIP 的 overall_recommended_K 为 2，且 2 个 cluster 均非空。两个样本都稳定地产生一个大型 common-appraisal-like 类别，而小型情绪相关类别的解释在两样本间较不稳定。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "报告原文献的值的评级分布、推论的一致情况，整理成表格，如下表所示："),
        "表 5 和表 6 只统计 Route A。结果数量 N=6，仅统计可以与原文正文直接配对的数值：2 个样本量，以及 2 个行为效应各自的 t 值和效应量。它们要么完全一致，要么属于舍入导致的微小偏差。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "* 结果数量N指在重复分析中，对重复分析结果与原结果进行配对比较的次数。对于每个目标效应，结果包括一组数值，如汇总效应估计(summary estimate，如t值/ F值)、置信区间界限(confidence interval bound)、效应量(effect size)样本大小(size effect)等，应将原文中报告的每个数值与重复结果进行比较。例如，在一个t检验中，原文献报告了t值、95%置信区间、cohen’s d和样本大小，则这个效应中N＝4。将各效应的N求和即为全体数量。"),
        "四类脑系统图不被纳入该 PE 统计，因为原文发布的是二值空间图而非单一统计量。对这些结果的复现评价单独使用共识图比较与跨样本稳定性指标。",
        size=10.5,
    )
    set_paragraph_text(
        find_paragraph(doc, "* 推论数量N指在重复分析中，对效应做出统计推断的次数。例如，仅进行了一个t检验，则N＝1；如果进行了一个2*2的方差分析，并进行简单效应分析，则有可能有7个统计推断：两个主效应的推论，一个交互作用的推论，四个可能的简单效应分析的推论，因此N＝7。如果报告的p值相对于重复的p值落在显著性水平边界的另一侧，则被归类为推论不一致；反之为推论一致。"),
        "Route A 的推论数量 N=6，包括两个关键行为推论和四类系统是否成功识别并形成共识图的推论；六项判断全部与原文方向一致。",
        size=10.5,
    )
    inconsistency_paragraphs = [p for p in doc.paragraphs if p.text.strip() == "若出现不一致的情况，需要文字总结出现了哪些不一致。"]
    set_paragraph_text(
        inconsistency_paragraphs[0],
        "原分析方法下未见推论不一致。需要强调的是，空间 Dice 低于 1 并不等于结论不一致，因为四类系统都被成功识别，且共识图几乎完整覆盖原文图。",
        size=10.5,
    )
    set_paragraph_text(
        find_paragraph(doc, "报告采用新方法后，推论与原文献推论的一致情况，整理成表格，如下表所示"),
        "表 7 只对应 Route B。Route B 评价的是“数据驱动类别结构是否支持原文四系统框架”，因此其一致性标准更严格：只有当后验结构稳定支持原文理论分化时，才记为一致。",
        size=11,
    )
    set_paragraph_text(
        inconsistency_paragraphs[1],
        "创新方法下出现两类不完全一致：一是 small emotion-related 类别在两样本中的标签解释不同，只能算部分对应原文；二是 K=2 至 K=6 的后验比较没有恢复四个稳定且非空的终端 cluster。这说明原文四系统更像是基于理论规则定义的功能分区，而不是后验自动聚类会自然给出的四类结构。",
        size=10.5,
    )
    set_paragraph_text(
        find_paragraph(doc, "结合下表，对原文献进行分析，推测可能导致可复现性检验结果差异的原因。对于重要的原因，逐段进行展开说明。"),
        "整体来看，Route A 显示原文的核心行为与系统识别结论具有较强计算可复现性，但空间边界存在系统性外扩；Route B 则提示这些理论系统未必会作为四个稳定终端 cluster 自然恢复。最关键的差异来源不是样本变化，而是两个路线回答的问题不同：Route A 检验“能否按原文规则重算”，Route B 检验“后验自动分群是否支持同样结构”。此外，共识图构建中的平滑乘积与阈值化步骤，以及 MATLAB/CANlab 与 R 实现差异，也会带来体素级边界变化。",
        size=11,
    )
    set_paragraph_text(
        find_paragraph(doc, "可自由总结数据分析过程中除计算可复现性问题外的其他思考，也可以包括对本课的建议。"),
        "本次项目最重要的体会是：Route A 与 Route B 回答的是不同层次的问题。Route A 检验公开材料能否把原文的理论分类规则重算出来；Route B 检验 ROI 效应模式在后验上会不会自动分成若干稳定类别。前者复现成功，并不保证后者一定得到 4 类；后者没有恢复 4 类，也不能直接否定原文结论。对神经影像结果做计算可复现性检验时，必须把“数值复现”“空间复现”和“模型解释复现”区分开来。",
        size=11,
    )

    heading_rewrites = {
        "2.2 原研究方法简介": "2.2 路线 A 与路线 B 的数据分析方法",
        "2.3 重复思路说明": "2.3 路线 A 与路线 B 的重复思路",
        "表 2 描述性统计结果的比较": "表 2 Route A 描述性统计结果的比较",
        "3.2.1 使用与原文献相同方法的推断性统计": "3.2.1 Route A：使用与原文献相同方法的推断性统计",
        "表 3 推断性统计结果的比较(原文献方法)": "表 3 Route A 推断性统计结果的比较(原文献方法)",
        "3.2.2 使用与原文献不同方法的推断性统计": "3.2.2 Route B：使用与原文献不同方法的推断性统计",
        "表 4 推断性统计结果(创新方法)": "表 4 Route B 推断性统计结果(贝叶斯 ROI-MCMC)",
        "3.3.1 使用与原文献相同方法": "3.3.1 Route A：使用与原文献相同方法",
        "表 5  结果可复现性的评估表": "表 5 Route A 结果可复现性的评估表",
        "表 6  推论的一致性的评估表(原分析方法)": "表 6 Route A 推论一致性的评估表",
        "3.3.2 使用与原文献不同方法": "3.3.2 Route B：使用与原文献不同方法",
        "表 7  推论的一致性的评估表(创新方法)": "表 7 Route B 推论一致性的评估表",
    }
    for old, new in heading_rewrites.items():
        set_paragraph_text(find_paragraph(doc, old), new, size=11.5, bold=True)

    reference_heading = find_paragraph(doc, "参考文献(APA格式)")
    set_paragraph_text(reference_heading, "参考文献(APA格式)", size=12, bold=True)
    reference_heading_index = find_paragraph_index(doc, "参考文献(APA格式)")
    set_paragraph_text(
        doc.paragraphs[reference_heading_index + 1],
        "Bo, K., Kraynak, T. E., Kwon, M., Sun, M., Gianaros, P. J., & Wager, T. D. (2024). A systems identification approach using Bayes factors to deconstruct the brain bases of emotion regulation. Nature Neuroscience. https://doi.org/10.1038/s41593-024-01605-7",
        size=10.5,
    )
    set_paragraph_text(
        doc.paragraphs[reference_heading_index + 2],
        "Kitzes, J., Turek, D., & Deniz, F. (2017). The practice of reproducible research: Case studies and lessons from the data-intensive sciences. University of California Press.",
        size=10.5,
    )


def validate_output(doc):
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    full_text = paragraph_text + "\n" + table_text
    placeholders = [
        "(APA 格式的参考文献)",
        "(osf 或 github 链接)",
        "变量一",
        "假设一",
        "（略）",
        "研究n",
        "待填写",
    ]
    remaining = [item for item in placeholders if item in full_text]
    if remaining:
        raise ValueError(f"Unfilled placeholders remain: {remaining}")


def save_with_fallback(doc):
    candidates = [
        OUT,
        OUT.parent / f"{OUT.stem}_修正版{OUT.suffix}",
        OUT.parent / f"{OUT.stem}_修正版_2{OUT.suffix}",
    ]
    last_error = None
    for path in candidates:
        try:
            doc.save(path)
            return path
        except PermissionError as exc:
            last_error = exc
    raise last_error


def build_report():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE)
    prune_before_title(doc)
    fill_paragraphs(doc)

    tables = doc.tables
    if len(tables) != 9:
        raise ValueError(f"Unexpected appendix table count: {len(tables)}")

    fill_group_table(tables[0])
    fill_literature_table(tables[1])
    fill_descriptive_table(tables[2])
    fill_original_inference_table(tables[3])
    fill_innovation_table(tables[4])
    fill_reproducibility_tables(tables[5], tables[6], tables[7])
    fill_reason_table(tables[8])

    validate_output(doc)
    return save_with_fallback(doc)


if __name__ == "__main__":
    print(build_report())
